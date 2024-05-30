import moviepy.editor as mp
import speech_recognition as sr
from pydub import AudioSegment
import os 
from pydub.silence import split_on_silence
import docx

# create a speech recognition object
r = sr.Recognizer()

# create a word document and add a title
mydoc = docx.Document()
mydoc.add_heading("Transcripción de audio", level=0)
mydoc.add_heading("Texto por fragmentos", level=1)
mydoc.add_paragraph("")

# convert mp4 to mp3 or wav direct in one function
def mp4tomp3 (video_path, audio_path):
    video = mp.VideoFileClip(video_path)
    video.audio.write_audiofile(audio_path)

'''
def mp3towav (audio_path, wav_path):
    sound = AudioSegment.from_mp3(audio_path)
    sound.export(wav_path, format="wav")
'''
# convert wav to text
def wavToText (audio_path):
    #recognizer = sr.Recognizer()
    #audio = sr.AudioFile(audio_path)

    #with audio as source:
    with sr.AudioFile(audio_path) as source:
        #audio_file = recognizer.record(source)
        audio_file = r.record(source)
    return r.recognize_google(audio_file, language="es-MX")
    #return text

# split the audio file into chunks, call wavToText for each chunk and add the text to the word document
def wavToText2(path):
    """Splitting the large audio file into chunks
    and apply speech recognition on each of these chunks"""
    # open the audio file using pydub
    sound = AudioSegment.from_file(path)  
    # split audio sound where silence is 500 miliseconds or more and get chunks
    chunks = split_on_silence(sound,
        # experiment with this value for your target audio file
        min_silence_len = 500,
        # adjust this per requirement
        silence_thresh = sound.dBFS-14,
        # keep the silence for 1 second, adjustable as well
        keep_silence=500,
    )
    folder_name = "audio-chunks"
    # create a directory to store the audio chunks
    if not os.path.isdir(folder_name):
        os.mkdir(folder_name)
    whole_text = ""
    # process each chunk 
    for i, audio_chunk in enumerate(chunks, start=1):
        # export audio chunk and save it in
        # the `folder_name` directory.
        chunk_filename = os.path.join(folder_name, f"chunk{i}.wav")
        audio_chunk.export(chunk_filename, format="wav")
        # recognize the chunk
        try:
            text = wavToText(chunk_filename)
            #text = transcribe_audio(chunk_filename)
        except sr.UnknownValueError as e:
            print("Error:", str(e))
        else:
            text = f"{text.capitalize()}. "
            print(chunk_filename, ":", text)
            mydoc.add_paragraph(text)
            whole_text += text
    # return the text for all chunks detected
    mydoc.add_page_break()
    mydoc.add_heading("Texto completo", level=1)
    mydoc.add_paragraph(whole_text)
    return whole_text

    
#mp4tomp3("video.mp4", "audio.mp3")
#mp4tomp3("Directorio/video.mp4", "Directorio/video_nuevo.wav")
#mp3towav("Test1.mp3", "Test1.wav") OLD
print(wavToText2("Directorio/documento.wav")) # One way to call the function
#wavToText2("Directorio/documento.wav") # Another way to call the function 
# set the author of the document
core_properties = mydoc.core_properties
core_properties.author = core_properties.autor + " BADLA"
# save the document
mydoc.save("Directorio/documento.docx")
#print(wavToText2("Test1.wav"))